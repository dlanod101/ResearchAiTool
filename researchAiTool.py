import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

# Load env vars
load_dotenv()

# Configure Gemini
genai.configure(api_key="AIzaSyC-ePg2dcvG-HUea7RQO5DELUQAnj8yIzE")
model = genai.GenerativeModel("gemini-1.5-flash")

# --- Prompt Chain Functions ---

def step_1_generate_context(topic: str) -> str:
    prompt = f"""
    Provide a detailed context and background on the topic: "{topic}".
    Include history, relevance, and recent trends.
    """
    response = model.generate_content(prompt)
    return response.text.strip()

def step_2_generate_draft(context: str) -> str:
    prompt = f"""
    Using the following context, generate a detailed first draft of an article:

    Context:
    {context}

    Ensure the draft is well-organized and informative.
    """
    response = model.generate_content(prompt)
    return response.text.strip()

def step_3_refine_draft(draft: str) -> str:
    prompt = f"""
    Refine the following draft to improve grammar, structure, and flow.
    Make it more polished and professional.

    Draft:
    {draft}
    """
    response = model.generate_content(prompt)
    return response.text.strip()

def generate_docx(topic: str, context: str, draft: str, refined: str) -> BytesIO:
    doc = Document()
    doc.add_heading(f"Gemini Content: {topic}", 0)

    doc.add_heading("1. Context", level=1)
    doc.add_paragraph(context)

    doc.add_heading("2. Draft", level=1)
    doc.add_paragraph(draft)

    doc.add_heading("3. Refined Version", level=1)
    doc.add_paragraph(refined)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit UI ---

st.set_page_config(page_title="Gemini Prompt Chainer", layout="centered")
st.title("🪄 Gemini Prompt Chaining Tool")
st.caption("Generate high-quality content in stages: Context → Draft → Refine")

topic = st.text_input("📌 Enter a topic to generate content:", placeholder="e.g., Artificial Intelligence in Healthcare")

if st.button("Generate"):
    if not topic.strip():
        st.warning("Please enter a topic.")
    else:
        with st.spinner("Generating context..."):
            context = step_1_generate_context(topic)
        st.subheader("1️⃣ Generated Context")
        st.markdown(context, unsafe_allow_html=True)

        with st.spinner("Generating draft..."):
            draft = step_2_generate_draft(context)
        st.subheader("2️⃣ First Draft")
        st.markdown(draft, unsafe_allow_html=True)

        with st.spinner("Refining draft..."):
            refined = step_3_refine_draft(draft)
        st.subheader("3️⃣ Refined Draft")
        st.markdown(refined, unsafe_allow_html=True)

        st.success("✅ Content generation complete!")

        # Offer DOCX download
        doc_buffer = generate_docx(topic, context, draft, refined)
        st.download_button(
            label="📄 Download Final Document (.docx)",
            data=doc_buffer,
            file_name=f"{topic.lower().replace(' ', '_')}_gemini_output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
